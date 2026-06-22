from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "One_Stop_Personalized_Career_and_Educational_Advisor_Report.docx"
FIG_DIR = ROOT / "generated_report_assets"
FIG_DIR.mkdir(exist_ok=True)


BLUE = RGBColor(31, 78, 121)
TEAL = RGBColor(0, 112, 96)
GRAY = RGBColor(90, 90, 90)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) < 18 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "D9EAF7")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
    style_table(table)
    doc.add_paragraph()
    return table


def add_title_page(doc):
    def centered(text, size=14, bold=False, space=6):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        r.bold = bold
        p.paragraph_format.space_after = Pt(space)
        return p

    centered("A Mini Project Report", 16, True)
    centered("on", 13)
    centered("ONE STOP PERSONALIZED CAREER AND EDUCATIONAL ADVISOR", 18, True, 18)
    centered("Submitted in partial fulfilment of the requirements for the award of the degree of", 12)
    centered("Bachelor of Engineering", 14, True)
    centered("in", 12)
    centered("Computer Science and Engineering", 14, True, 20)
    centered("By", 13, True)
    centered("A. PRAVEEN                          1608-23-733-198", 12, True)
    centered("R. PRASANNA KUMAR       1608-23-733-234", 12, True)
    centered("V. SAI VARDHAN REDDY    1608-23-733-225", 12, True, 20)
    centered("Under the guidance of", 12)
    centered("Mr. T HARIKRISHNA", 13, True)
    centered("Assistant Professor", 12, False, 24)
    centered("DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING", 13, True)
    centered("Matrusri Engineering College", 13, True)
    centered("(Affiliated to Osmania University, Approved by AICTE)", 11)
    centered("Saidabad, Hyderabad - 500059", 11)
    centered("(2025-2026)", 11)
    doc.add_page_break()


def add_front_matter(doc):
    add_heading(doc, "CERTIFICATE", 1, center=True)
    add_para(doc, 'This is to certify that a mini project report entitled "ONE STOP PERSONALIZED CAREER AND EDUCATIONAL ADVISOR" is being submitted by A. Praveen, R. Prasanna Kumar and V. Sai Vardhan Reddy in partial fulfilment of the requirements for the award of Bachelor of Engineering in Computer Science and Engineering. The work has been carried out under the guidance of Mr. T. Harikrishna, Assistant Professor, Department of Computer Science and Engineering, Matrusri Engineering College.')
    doc.add_paragraph("\n\nProject Coordinator                         Project Guide                         H.O.D")
    doc.add_paragraph("Mrs. M. Priyanka                           Mr. T. Harikrishna                    Dr. T. Raghunadha Reddy")
    doc.add_paragraph("Assistant Professor                         Assistant Professor                   Professor & Head")
    doc.add_paragraph("External Examiner(s): ________________________________")
    doc.add_page_break()

    add_heading(doc, "DECLARATION", 1, center=True)
    add_para(doc, "We hereby declare that the mini project entitled One Stop Personalized Career and Educational Advisor is a record of original work carried out by us under the guidance of Mr. T. Harikrishna, Assistant Professor, Department of Computer Science and Engineering, Matrusri Engineering College. The project has not been submitted earlier to any other university or institution for the award of any degree or diploma.")
    doc.add_paragraph("\nA PRAVEEN (1608-23-733-198)                               _______________________")
    doc.add_paragraph("R PRASANNA KUMAR (1608-23-733-234)              _______________________")
    doc.add_paragraph("V SAI VARDHAN REDDY (1608-23-733-225)         _______________________")
    doc.add_page_break()

    add_heading(doc, "ACKNOWLEDGEMENT", 1, center=True)
    for text in [
        "It is our privilege to express our sincere gratitude to our guide Mr. T. Harikrishna, Assistant Professor, Department of Computer Science and Engineering, for his valuable guidance and encouragement throughout the project.",
        "We express our sincere thanks to the mini project coordinator Mrs. M. Priyanka for providing continuous support, feedback and motivation during the development of this application.",
        "We are thankful to Dr. T. Raghunadha Reddy, Professor and Head, Department of Computer Science and Engineering, Matrusri Engineering College, for providing the required academic environment and facilities.",
        "We also thank all teaching and non-teaching staff members, our classmates, friends and family members for their support in completing this project successfully."
    ]:
        add_para(doc, text)
    doc.add_page_break()

    add_heading(doc, "ABSTRACT", 1, center=True)
    abstract = [
        "Career planning and education selection are complex problems for students because the decision depends on academic strengths, technical skills, interests, aptitude, resume quality, college preferences and job-readiness. Many students either depend on generic advice or choose paths without enough evidence about their own profile.",
        "The One Stop Personalized Career and Educational Advisor is a MERN stack web application designed to solve this problem by combining authentication, profile-based career prediction, resume NLP analysis, college recommendation, job navigation, personalized roadmaps, skill quizzes, expert guidance and progress tracking in one platform.",
        "The system uses React.js for a responsive dashboard, Node.js and Express.js for REST APIs, MongoDB for persistent storage, and Python-based machine learning services for career prediction and similar-profile matching. The career prediction module uses Random Forest logic over structured profile features, while KNN is used to identify similar user profiles and display community feedback. Resume analysis uses NLP-style extraction of skills, education, projects and skill gaps for a selected target role.",
        "The final system provides students with career recommendations, learning resources, YouTube links, free certifications, job-site navigation links, college matching, roadmap progress tracking and feedback storage. The project demonstrates a practical full-stack implementation of a personalized career advisory platform."
    ]
    for text in abstract:
        add_para(doc, text)
    doc.add_page_break()


def add_heading(doc, text, level=1, center=False):
    p = doc.add_heading(text, level=level)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = BLUE if level == 1 else TEAL
    return p


def add_para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)


def draw_box_diagram(name, title, boxes, arrows):
    img = Image.new("RGB", (1400, 850), "white")
    draw = ImageDraw.Draw(img)
    try:
        font_title = ImageFont.truetype("arialbd.ttf", 38)
        font = ImageFont.truetype("arial.ttf", 24)
        font_bold = ImageFont.truetype("arialbd.ttf", 24)
    except Exception:
        font_title = font = font_bold = ImageFont.load_default()
    draw.rectangle((0, 0, 1400, 80), fill=(31, 78, 121))
    draw.text((40, 22), title, fill="white", font=font_title)
    for key, (x, y, w, h, label, fill) in boxes.items():
        draw.rounded_rectangle((x, y, x + w, y + h), radius=18, fill=fill, outline=(50, 70, 95), width=3)
        lines = label.split("\n")
        for i, line in enumerate(lines):
            draw.text((x + 20, y + 20 + i * 30), line, fill=(15, 23, 42), font=font_bold if i == 0 else font)
    for a, b, label in arrows:
        ax, ay, aw, ah, *_ = boxes[a]
        bx, by, bw, bh, *_ = boxes[b]
        start = (ax + aw, ay + ah // 2) if bx > ax else (ax + aw // 2, ay + ah)
        end = (bx, by + bh // 2) if bx > ax else (bx + bw // 2, by)
        draw.line((start, end), fill=(31, 78, 121), width=4)
        # arrow head
        ex, ey = end
        draw.polygon([(ex, ey), (ex - 14, ey - 8), (ex - 14, ey + 8)], fill=(31, 78, 121))
        if label:
            mx, my = (start[0] + end[0]) // 2, (start[1] + end[1]) // 2
            draw.text((mx - 40, my - 28), label, fill=(90, 90, 90), font=font)
    path = FIG_DIR / f"{name}.png"
    img.save(path)
    return path


def make_diagrams():
    diagrams = []
    diagrams.append(draw_box_diagram(
        "architecture", "System Architecture",
        {
            "u": (40, 150, 250, 120, "Student/User\nBrowser", (219, 234, 254)),
            "r": (360, 125, 300, 170, "React Frontend\nDashboard, Tests,\nResume, Guidance", (224, 242, 254)),
            "e": (740, 125, 280, 170, "Express API\nAuth, Advisor,\nJobs, Resume", (220, 252, 231)),
            "m": (1100, 120, 240, 120, "MongoDB\nUsers, Tests,\nFeedback", (254, 243, 199)),
            "ml": (740, 420, 280, 150, "Python ML Service\nRandom Forest,\nKNN", (237, 233, 254)),
            "ext": (1100, 420, 240, 150, "External Links\nJobs, YouTube,\nCertifications", (255, 237, 213)),
        },
        [("u", "r", ""), ("r", "e", "REST"), ("e", "m", ""), ("e", "ml", ""), ("e", "ext", "links")]
    ))
    diagrams.append(draw_box_diagram(
        "use_case", "Use Case Diagram",
        {
            "stu": (40, 350, 200, 100, "Student", (219, 234, 254)),
            "auth": (370, 110, 300, 90, "Login/Register", (224, 242, 254)),
            "test": (370, 230, 300, 90, "Career Profiling Test", (224, 242, 254)),
            "resume": (370, 350, 300, 90, "Resume NLP Analysis", (224, 242, 254)),
            "job": (370, 470, 300, 90, "Job Recommendations", (224, 242, 254)),
            "road": (760, 230, 300, 90, "Guidance Roadmap", (220, 252, 231)),
            "feed": (760, 350, 300, 90, "Share Feedback", (220, 252, 231)),
            "college": (760, 470, 300, 90, "College Matching", (220, 252, 231)),
        },
        [("stu", "auth", ""), ("stu", "test", ""), ("stu", "resume", ""), ("stu", "job", ""), ("test", "road", ""), ("road", "feed", ""), ("stu", "college", "")]
    ))
    diagrams.append(draw_box_diagram(
        "class", "Class Diagram",
        {
            "user": (50, 130, 290, 170, "User\n_id, name, email\npasswordHash, provider", (219, 234, 254)),
            "profile": (420, 130, 330, 170, "UserProfile\nskills, targetRole\nroadmapProgress", (224, 242, 254)),
            "test": (820, 130, 320, 170, "TestResult\nanswers, scores\nrecommendation", (220, 252, 231)),
            "resume": (420, 430, 330, 170, "ResumeAnalysis\ntargetRole, skills\nmissingSkills, ATS", (254, 243, 199)),
            "feedback": (820, 430, 320, 170, "RecommendationFeedback\nselectedCareer\noutcome, notes", (255, 237, 213)),
        },
        [("user", "profile", "1:1"), ("profile", "test", "1:N"), ("profile", "resume", "1:N"), ("test", "feedback", "1:N")]
    ))
    diagrams.append(draw_box_diagram(
        "sequence", "Career Prediction Sequence",
        {
            "user": (40, 150, 230, 100, "User", (219, 234, 254)),
            "react": (330, 150, 230, 100, "React UI", (224, 242, 254)),
            "api": (620, 150, 230, 100, "Express API", (220, 252, 231)),
            "ml": (910, 150, 230, 100, "ML Service", (237, 233, 254)),
            "db": (620, 430, 230, 100, "MongoDB", (254, 243, 199)),
        },
        [("user", "react", "answers"), ("react", "api", "scores"), ("api", "ml", "predict"), ("api", "db", "save"), ("api", "react", "result")]
    ))
    diagrams.append(draw_box_diagram(
        "activity", "Guidance and Feedback Activity",
        {
            "start": (50, 140, 240, 90, "Start", (219, 234, 254)),
            "profile": (360, 140, 270, 90, "Load saved profile", (224, 242, 254)),
            "road": (700, 140, 280, 90, "Show roadmap", (220, 252, 231)),
            "complete": (700, 330, 280, 90, "Mark steps complete", (254, 243, 199)),
            "feedback": (1030, 330, 260, 90, "Store feedback", (255, 237, 213)),
        },
        [("start", "profile", ""), ("profile", "road", ""), ("road", "complete", ""), ("complete", "feedback", "100%")]
    ))
    diagrams.append(draw_box_diagram(
        "deployment", "Deployment Diagram",
        {
            "browser": (60, 170, 270, 110, "Client Browser\nReact/Vite", (219, 234, 254)),
            "node": (430, 160, 280, 130, "Node Server\nExpress APIs", (220, 252, 231)),
            "python": (820, 160, 280, 130, "Python Service\nML/NLP", (237, 233, 254)),
            "mongo": (430, 430, 280, 120, "MongoDB Server\nPersistent Data", (254, 243, 199)),
        },
        [("browser", "node", "HTTP"), ("node", "python", "REST"), ("node", "mongo", "Mongoose")]
    ))
    diagrams.append(draw_box_diagram(
        "er", "Database Relationship Diagram",
        {
            "user": (70, 130, 250, 120, "Users", (219, 234, 254)),
            "profile": (410, 130, 250, 120, "UserProfiles", (224, 242, 254)),
            "tests": (750, 130, 250, 120, "TestResults", (220, 252, 231)),
            "resume": (410, 390, 250, 120, "ResumeAnalyses", (254, 243, 199)),
            "feedback": (750, 390, 250, 120, "Feedback", (255, 237, 213)),
        },
        [("user", "profile", "1:1"), ("user", "tests", "1:N"), ("user", "resume", "1:N"), ("user", "feedback", "1:N")]
    ))
    diagrams.append(draw_box_diagram(
        "module", "Module Interaction Diagram",
        {
            "dash": (50, 160, 270, 110, "Dashboard", (219, 234, 254)),
            "career": (400, 90, 270, 110, "Career Prediction", (224, 242, 254)),
            "resume": (400, 250, 270, 110, "Resume NLP", (220, 252, 231)),
            "college": (400, 410, 270, 110, "College Match", (254, 243, 199)),
            "guidance": (760, 250, 270, 110, "Guidance Path", (255, 237, 213)),
        },
        [("dash", "career", ""), ("dash", "resume", ""), ("dash", "college", ""), ("career", "guidance", ""), ("resume", "guidance", "")]
    ))
    return diagrams


def add_figure(doc, path, caption):
    doc.add_picture(str(path), width=Inches(6.2))
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)
        r.italic = True


def setup_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Times New Roman"
        styles[name].font.color.rgb = BLUE if name == "Heading 1" else TEAL
    return doc


def add_contents_tables(doc):
    add_heading(doc, "TABLE OF CONTENTS", 1, center=True)
    toc_rows = [
        ("", "Abstract", "v"), ("", "Table of Contents", "vi"), ("", "List of Figures", "vii"), ("", "List of Tables", "viii"),
        ("Chapter 1", "Introduction", "1"), ("Chapter 2", "Requirement Analysis & SRS", "8"), ("Chapter 3", "System Design", "14"),
        ("Chapter 4", "Implementation, Testing and Maintenance", "26"), ("Chapter 5", "Results and Discussions", "39"),
        ("Chapter 6", "Conclusion and Future Scope", "48"), ("", "References/Bibliography", "52"),
        ("1.1", "Introduction to Project", "1"), ("1.2", "Project Category", "2"), ("1.3", "Objectives", "3"),
        ("1.4", "Scope of the Problem", "4"), ("1.5", "Identification of Need", "5"), ("1.6", "Existing System", "5"),
        ("1.7", "Limitations of Existing System", "6"), ("1.8", "Proposed System", "6"), ("1.9", "Unique Features", "7"),
        ("2.1", "Feasibility Study", "8"), ("2.2", "Software Requirement Specification", "9"), ("2.3", "Validation", "11"),
        ("2.4", "Expected Hurdles", "12"), ("2.5", "SDLC Model", "13"), ("3.1", "Design Approach", "14"),
        ("3.2", "System Architecture", "15"), ("3.3", "UML Diagrams", "16"), ("3.4", "Interface Relationship", "22"),
        ("3.5", "Database Design", "23"), ("3.6", "User Interface Design", "24"), ("3.7", "REST API Endpoints", "25"),
        ("4.1", "Technologies Used", "26"), ("4.2", "Coding Standards", "30"), ("4.3", "Testing Techniques", "34"),
        ("4.4", "Executable Code Listings", "36"), ("5.1", "User Interface Representation", "39"), ("5.2", "Various Modules", "41"),
        ("5.3", "System Snapshots", "44"), ("5.4", "Detailed Test Cases", "46"), ("6.1", "Conclusion", "48"),
        ("6.2", "Future Scope", "49")
    ]
    add_table(doc, ["Chapter / Section", "Title", "Page No."], toc_rows)
    doc.add_page_break()
    add_heading(doc, "LIST OF FIGURES", 1, center=True)
    fig_rows = [
        ("Figure 3.1", "System Architecture Diagram", "15"), ("Figure 3.2", "Use Case Diagram", "16"),
        ("Figure 3.3", "Class Diagram", "17"), ("Figure 3.4", "Sequence Diagram", "18"),
        ("Figure 3.5", "Activity Diagram", "19"), ("Figure 3.6", "Deployment Diagram", "20"),
        ("Figure 3.7", "Database Relationship Diagram", "21"), ("Figure 3.8", "Module Interaction Diagram", "22"),
        ("Figure 5.1", "Login and Dashboard Screen", "40"), ("Figure 5.2", "Career Prediction Result Dialog", "41"),
        ("Figure 5.3", "Resume NLP Skill Gap Dashboard", "42"), ("Figure 5.4", "Guidance Roadmap and Resources", "43"),
        ("Figure 5.5", "Job Finder with External Links", "44"), ("Figure 5.6", "College Recommendation Screen", "45"),
        ("Figure 5.7", "Skill Quiz Section", "46")
    ]
    add_table(doc, ["Figure No.", "Figure Title", "Page No."], fig_rows)
    doc.add_page_break()
    add_heading(doc, "LIST OF TABLES", 1, center=True)
    table_rows = [
        ("Table 3.1", "MongoDB Collections Data Dictionary", "23"),
        ("Table 3.2", "REST API Endpoints", "25"),
        ("Table 4.1", "Testing and Validation Test Cases", "43"),
        ("Table 2.1", "Functional Requirements", "10"),
        ("Table 2.2", "Non-Functional Requirements", "11"),
        ("Table 3.3", "User Profile Fields", "23"),
        ("Table 3.4", "Resume Analysis Fields", "24"),
        ("Table 3.5", "Feedback Analytics Fields", "24"),
        ("Table 5.1", "Module-wise Output Summary", "42"),
        ("Table 5.2", "Future Enhancement Plan", "50")
    ]
    add_table(doc, ["Table No.", "Table Title", "Page No."], table_rows)
    doc.add_page_break()


def add_chapters(doc, diagrams):
    add_heading(doc, "Chapter 1: Introduction", 1)
    sections = [
        ("1.1 Introduction to Project", [
            "The One Stop Personalized Career and Educational Advisor is a full-stack web application that helps students make informed academic and career decisions. It combines career profiling tests, manual skill input, resume analysis, college recommendation, job navigation, guidance roadmaps and progress tracking into a unified dashboard.",
            "The application is designed for students who are unsure about which career path, skill track, college option or job role fits their current profile. Instead of giving generic advice, the system collects structured user data and converts it into actionable recommendations."
        ]),
        ("1.2 Project Category", [
            "The project belongs to the category of AI-assisted web-based educational technology and career advisory systems. It is implemented using the MERN stack with an additional Python ML service for prediction and similarity-based recommendations."
        ]),
        ("1.3 Objectives", [
            "The primary objective is to provide a personalized career and education advisory platform that can evaluate a student profile and recommend suitable career paths, skills, courses, colleges and jobs."
        ]),
        ("1.4 Scope of the Problem", [
            "Students often face confusion while choosing between software engineering, data science, AI, cybersecurity, cloud, full-stack development, business analysis and other roles. The system focuses on reducing this uncertainty through structured assessment and continuous guidance."
        ]),
        ("1.5 Identification of Need", [
            "There is a need for a centralized platform that connects career prediction, resume analysis, skill development, college selection and job exploration. Existing guidance is usually fragmented across different websites and does not preserve user progress."
        ]),
        ("1.6 Existing System", [
            "Traditional career guidance usually depends on counselling sessions, manual aptitude tests or generic online articles. Resume checking, job search and college comparison are handled on separate platforms."
        ]),
        ("1.7 Limitations of Existing System", [
            "Existing systems often lack personalization, feedback storage, similar-profile comparison, roadmap tracking and integration between education and job-readiness modules."
        ]),
        ("1.8 Proposed System", [
            "The proposed system provides login-based personalized guidance. It stores user profiles in MongoDB, predicts careers from test/manual inputs, analyzes resumes using NLP logic, tracks roadmap progress and stores feedback for future similar-user recommendations."
        ]),
        ("1.9 Unique Features of the System", [
            "The platform includes Random Forest career prediction, KNN similar-profile matching, resume NLP analysis, dynamic skill quizzes, roadmap progress tracking, feedback-based learning, college recommendations and real job-site navigation links."
        ])
    ]
    for heading, paras in sections:
        add_heading(doc, heading, 2)
        for text in paras:
            add_para(doc, text)
    add_bullets(doc, [
        "Unified MERN dashboard for career, education, resume and jobs.",
        "Google OAuth and normal login/register authentication.",
        "Stored user progress and feedback in MongoDB.",
        "Guidance path with YouTube courses, certifications and learning websites.",
        "Profile-matched job links using Naukri, LinkedIn, Indeed and Internshala."
    ])
    doc.add_page_break()

    add_heading(doc, "Chapter 2: Requirement Analysis & SRS", 1)
    add_heading(doc, "2.1 Feasibility Study", 2)
    add_para(doc, "The project is technically feasible because it uses open-source and widely supported technologies. React.js is used for the frontend, Express.js for REST APIs, MongoDB for data persistence and Python services for ML/NLP operations.")
    add_para(doc, "Economically, the prototype is feasible because it can run locally using free tools. Google OAuth can be used for development without payment, and MongoDB Community Server is sufficient for academic demonstration.")
    add_para(doc, "Operationally, the system is feasible because students interact with clear dashboard options such as career profiling, manual skills, resume NLP, college recommendation, job finder, guidance and progress tracking.")
    add_heading(doc, "2.2 Software Requirement Specification (SRS)", 2)
    add_bullets(doc, [
        "Authentication: Users must register, login and optionally use Google OAuth.",
        "Career Prediction: The system must accept profile features and generate career recommendations.",
        "Resume NLP: The system must accept text/PDF/image resume input and identify skills, gaps and ATS score.",
        "Guidance: The system must generate learning paths, resources, certifications and roadmap progress.",
        "Feedback: The system must store selected career, outcome and notes for similar-profile recommendations.",
        "Job Finder: The system must match user profile to job roles and provide job-site navigation links.",
        "College Recommendation: The system must recommend colleges based on exam type, rank/percentile, budget, location and course."
    ])
    add_heading(doc, "2.3 Validation", 2)
    add_para(doc, "Validation is performed by checking authentication, protected routes, prediction output, resume parsing, feedback storage, roadmap progress saving and job link generation.")
    add_heading(doc, "2.4 Expected Hurdles", 2)
    add_para(doc, "Expected hurdles include limited real admission datasets, external job-site API restrictions, resume format variations and the need to balance prototype reliability with advanced ML features.")
    add_heading(doc, "2.5 SDLC Model", 2)
    add_para(doc, "The project follows an iterative prototype model. Core modules were built first, then improved through repeated testing, user feedback and UI refinements.")
    doc.add_page_break()

    add_heading(doc, "Chapter 3: System Design", 1)
    add_heading(doc, "3.1 Design Approach", 2)
    add_para(doc, "The system follows modular full-stack design. React components represent dashboard modules, Express routes expose APIs, MongoDB models store data and Python services handle ML predictions.")
    add_heading(doc, "3.2 System Architecture", 2)
    add_para(doc, "The architecture contains a React client, Express API server, MongoDB database, Python ML service and external navigation resources for jobs, videos and certifications.")
    captions = [
        "Figure 3.1 System Architecture Diagram", "Figure 3.2 Use Case Diagram", "Figure 3.3 Class Diagram",
        "Figure 3.4 Sequence Diagram", "Figure 3.5 Activity Diagram", "Figure 3.6 Deployment Diagram",
        "Figure 3.7 Database Relationship Diagram", "Figure 3.8 Module Interaction Diagram"
    ]
    add_heading(doc, "3.3 UML Diagrams", 2)
    for path, caption in zip(diagrams, captions):
        add_figure(doc, path, caption)
    add_heading(doc, "3.4 Interface Relationship & Dependencies", 2)
    add_para(doc, "The frontend depends on REST API responses for profile, recommendations, resumes, college matches, jobs and feedback. The backend depends on MongoDB models and optionally the Python ML service.")
    add_heading(doc, "3.5 Database Design", 2)
    rows = [
        ("_id", "ObjectId", "YES", "Primary identifier for each document."),
        ("User.name", "String", "YES", "Full name of the registered student."),
        ("User.email", "String", "YES", "Unique login credential."),
        ("User.passwordHash", "String", "NO", "Hashed password for normal login."),
        ("User.provider", "String", "NO", "OAuth provider such as Google."),
        ("User.avatarUrl", "String", "NO", "Profile image from OAuth account."),
        ("UserProfile.user", "ObjectId", "YES", "Reference to the authenticated user."),
        ("UserProfile.skills", "Array", "NO", "Current skills entered by the user."),
        ("UserProfile.targetRole", "String", "NO", "Saved or confirmed career path."),
        ("UserProfile.completedSkills", "Array", "NO", "Completed roadmap or skill items."),
        ("UserProfile.roadmapProgress", "Number", "NO", "Percentage of guidance roadmap completed."),
        ("TestResult.answers", "Map", "YES", "Career profiling or manual input answers."),
        ("TestResult.scores", "Map", "YES", "Model feature scores derived from inputs."),
        ("TestResult.recommendation", "Object", "NO", "Career, confidence, alternatives and learning path."),
        ("ResumeAnalysis.targetRole", "String", "YES", "Role selected by the user for resume matching."),
        ("ResumeAnalysis.extractedSkills", "Array", "NO", "Skills detected from resume text."),
        ("ResumeAnalysis.missingSkills", "Array", "NO", "Role-specific missing skills."),
        ("ResumeAnalysis.atsScore", "Number", "NO", "Resume fit/ATS-style score."),
        ("RecommendationFeedback.selectedCareer", "String", "YES", "Career selected by the user."),
        ("RecommendationFeedback.userFeatures", "Map", "NO", "Feature vector used for KNN comparison."),
        ("RecommendationFeedback.outcome", "String", "YES", "Job, internship, learning or not interested status."),
        ("RecommendationFeedback.feedbackNotes", "String", "NO", "User comments after following recommendation."),
        ("CareerAnalytics.career", "String", "YES", "Career name used for aggregate analytics."),
        ("CareerAnalytics.successRate", "Number", "NO", "Calculated success rate from feedback.")
    ]
    add_table(doc, ["Field", "Type", "Required", "Description"], rows)
    add_heading(doc, "3.6 User Interface Design", 2)
    add_para(doc, "The interface is a dashboard-oriented React UI with sidebar navigation, cards, forms, dialogs, flowcharts, progress states and responsive layouts.")
    add_heading(doc, "3.7 REST API Endpoints", 2)
    api_rows = [
        ("/api/auth/register", "POST", "Creates user account", "No"),
        ("/api/auth/login", "POST", "Authenticates user and returns token", "No"),
        ("/api/auth/google", "GET", "Starts Google OAuth login", "No"),
        ("/api/advisor/profile", "GET/PUT", "Reads or updates user profile", "JWT"),
        ("/api/advisor/predict-profile", "POST", "Predicts career from 21+ features", "JWT"),
        ("/api/resumes/analyze_resume", "POST", "Analyzes resume using NLP logic", "JWT"),
        ("/api/advisor/jobs", "GET", "Returns profile-matched job links", "JWT")
    ]
    add_table(doc, ["Endpoint Route", "Method", "Description", "Auth Required"], api_rows)
    doc.add_page_break()

    add_heading(doc, "Chapter 4: Implementation, Testing and Maintenance", 1)
    add_heading(doc, "4.1 Introduction to Languages, IDEs, Tools and Technologies used", 2)
    add_bullets(doc, [
        "Frontend: React.js, Vite, JavaScript, CSS and lucide-react icons.",
        "Backend: Node.js, Express.js, Passport.js, JWT, bcrypt and Mongoose.",
        "Database: MongoDB for users, profiles, tests, resumes, feedback and analytics.",
        "ML/NLP: Python service using Random Forest-style prediction, KNN similarity and NLP keyword extraction.",
        "Tools: VS Code, Postman, browser developer tools and npm scripts."
    ])
    add_heading(doc, "4.2 Coding standards", 2)
    add_para(doc, "The code is organized into frontend components, backend routes, MongoDB models and ML service files. Reusable API helpers, protected routes and modular panels improve maintainability.")
    add_heading(doc, "4.3 Testing Techniques and Test Plans", 2)
    add_para(doc, "Testing includes login tests, API tests through Postman, frontend build verification, backend syntax checks, prediction testing, resume analysis testing and user-flow testing across dashboard modules.")
    test_rows = [
        ("AUTH-01", "User login", "Valid credentials", "Dashboard opens", "PASS"),
        ("AUTH-02", "Google OAuth", "Google account", "User session created", "PASS"),
        ("CP-01", "Career profiling", "Completed MCQ test", "Career recommendations shown", "PASS"),
        ("MS-01", "Manual skills", "Feature inputs", "Random Forest recommendations shown", "PASS"),
        ("KNN-01", "Similar profiles", "Stored feedback from another user", "Similar feedback displayed", "PASS"),
        ("RES-01", "Resume NLP", "Resume text/file and target role", "ATS score and skill gaps shown", "PASS"),
        ("GUIDE-01", "Roadmap progress", "Mark roadmap step complete", "Progress saved in profile", "PASS"),
        ("JOB-01", "Job finder", "Saved career profile", "Matched job cards and job-site links shown", "PASS")
    ]
    add_table(doc, ["Test ID", "Description", "Input", "Expected Output", "Status"], test_rows)
    add_heading(doc, "4.4 Executable Code Listings", 2)
    add_para(doc, "Important executable modules include server/src/routes/advisor.js for recommendation APIs, server/src/routes/auth.js for authentication, server/src/routes/resumes.js for resume NLP analysis, client/src/components/Dashboard.jsx for the main dashboard, client/src/components/ResumePanel.jsx for resume scanning and ml-service/train_model.py for model training.")
    doc.add_page_break()

    add_heading(doc, "Chapter 5: Results and Discussions", 1)
    add_heading(doc, "5.1 User Interface Representation", 2)
    add_para(doc, "The implemented dashboard provides a landing section, major action cards, sidebar navigation, protected user profile area and module-specific panels. The UI supports career profiling, manual skills, resume NLP, college recommendation, test your skills, job finder and guidance.")
    add_heading(doc, "5.2 Brief Description of Various Modules", 2)
    add_bullets(doc, [
        "Career Profiling Test: Converts MCQ answers into model features for prediction.",
        "Manual Skills: Accepts structured academic and skill inputs.",
        "Resume NLP: Extracts skills, role match, ATS score and missing keywords.",
        "Guidance Path: Shows roadmap, videos, certifications, websites and progress.",
        "Job Finder: Matches profile to roles and opens real job-site search links.",
        "College Recommendation: Uses entrance profile and preferences to rank colleges.",
        "Skill Quiz: Generates dynamic quizzes from a separate dataset file."
    ])
    add_heading(doc, "5.3 Snapshots of system with brief detail", 2)
    add_para(doc, "Screenshots may be inserted for the login page, dashboard, career profiling result dialog, resume NLP scanner, guidance roadmap, job finder and college recommendation screens.")
    add_heading(doc, "5.4 Detailed Test Cases", 2)
    add_para(doc, "The test cases confirm that all major user flows are connected and that data is stored for later sessions. The guidance page restores previous recommendations when an old user logs in.")
    doc.add_page_break()

    add_heading(doc, "Chapter 6: Conclusion and Future Scope", 1)
    add_heading(doc, "6.1 Conclusion", 2)
    add_para(doc, "The One Stop Personalized Career and Educational Advisor successfully demonstrates a working full-stack career guidance platform. It integrates user authentication, profile-based recommendations, resume NLP analysis, KNN feedback comparison, college matching, job navigation, roadmaps and progress tracking.")
    add_heading(doc, "6.2 Future Scope", 2)
    add_bullets(doc, [
        "Move quiz questions from local dataset files to MongoDB with admin management.",
        "Use larger official datasets for college cutoff prediction.",
        "Integrate verified job APIs where available.",
        "Improve resume NLP using transformer-based pretrained models.",
        "Deploy the system on cloud platforms with production-grade security.",
        "Add mentor appointment scheduling and video consultation support."
    ])
    doc.add_page_break()

    add_heading(doc, "References/Bibliography", 1)
    refs = [
        "MongoDB Documentation, https://www.mongodb.com/docs/",
        "Express.js Documentation, https://expressjs.com/",
        "React Documentation, https://react.dev/",
        "Node.js Documentation, https://nodejs.org/",
        "Scikit-learn Documentation, https://scikit-learn.org/",
        "Passport.js Documentation, https://www.passportjs.org/",
        "freeCodeCamp, Microsoft Learn, Kaggle Learn, Cisco Skills For All and AWS Skill Builder for learning resource references."
    ]
    for ref in refs:
        add_para(doc, ref)


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("One Stop Personalized Career and Educational Advisor")
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY


def main():
    diagrams = make_diagrams()
    doc = setup_doc()
    add_title_page(doc)
    add_front_matter(doc)
    add_contents_tables(doc)
    add_chapters(doc, diagrams)
    add_footer(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
